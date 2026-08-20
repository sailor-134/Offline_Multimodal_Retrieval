import os
import unittest
from docx import Document
from file_parser import ParserFactory

class TestDOCXParser(unittest.TestCase):
    def setUp(self):
        self.test_file = os.path.abspath("test_sample.docx")
        doc = Document()
        doc.add_paragraph("这是DOCX测试文本")
        doc.add_paragraph("用于验证DOCX解析功能")
        doc.save(self.test_file)

    def tearDown(self):
        if os.path.exists(self.test_file):
            os.remove(self.test_file)

    def test_normal_parse(self):
        parser = ParserFactory.get_parser(self.test_file)
        result = parser.parse(self.test_file)
        self.assertTrue(result.success)
        self.assertIn("测试文本", result.content)
        self.assertEqual(result.metadata["format"], ".docx")

    def test_file_not_exist(self):
        fake_path = os.path.abspath("not_exist.docx")
        parser = ParserFactory.get_parser(fake_path)
        result = parser.parse(fake_path)
        self.assertFalse(result.success)

if __name__ == "__main__":
    unittest.main()